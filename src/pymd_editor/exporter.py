from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import markdown2


class WordExporter:
    """Export Markdown to Word (.docx) format."""

    def export(self, markdown_text: str, output_path: Path) -> None:
        """Convert markdown to Word document."""
        doc = Document()
        
        # Parse markdown to get structure
        html = markdown2.markdown(markdown_text or "", extras=["fenced-code-blocks", "tables"])
        
        # Simple approach: convert markdown to plain paragraphs
        # For MVP, we'll just write the markdown text directly
        # In production, you'd parse HTML/markdown AST and map to Word styles
        
        # Add title
        title = doc.add_heading("Markdown Document", level=0)
        
        # Add content as paragraphs (simple split by lines)
        lines = markdown_text.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
        
        doc.save(str(output_path))


class PDFExporter:
    """Export Markdown to PDF format using Qt's native printing."""

    def export_via_webengine(self, web_view, output_path: Path) -> None:
        """
        Export PDF using QWebEngineView's native print capability.
        This method should be called with the preview QWebEngineView.
        """
        # This will be called from the main window with the preview widget
        from PyQt6.QtCore import QMarginsF
        from PyQt6.QtGui import QPageLayout, QPageSize
        
        page_layout = QPageLayout(
            QPageSize(QPageSize.PageSizeId.A4),
            QPageLayout.Orientation.Portrait,
            QMarginsF(15, 15, 15, 15)
        )
        
        # Use Qt's native PDF printing - no GTK dependencies
        web_view.page().printToPdf(str(output_path), page_layout)
    
    def export(self, markdown_text: str, output_path: Path, dark: bool = False) -> None:
        """
        Fallback method for direct markdown-to-PDF conversion.
        Note: This requires weasyprint and GTK dependencies.
        """
        try:
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
        except ImportError as e:
            raise ImportError(
                "weasyprint not installed. Install with: pip install weasyprint\n"
                "Note: On Windows, weasyprint may require GTK3 runtime. "
                "Recommended: Use 'Print Preview' export instead (no dependencies)."
            ) from e
        
        # Convert markdown to HTML
        html_body = markdown2.markdown(
            markdown_text or "",
            extras=["fenced-code-blocks", "tables", "strike", "task_list"]
        )
        
        css = DARK_PDF_CSS if dark else LIGHT_PDF_CSS
        
        html_doc = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8" />
<style>{css}</style>
</head>
<body>
<div class="content">{html_body}</div>
</body>
</html>
"""
        
        font_config = FontConfiguration()
        HTML(string=html_doc).write_pdf(str(output_path), font_config=font_config)


LIGHT_PDF_CSS = """
@page { size: A4; margin: 2cm; }
body { 
    font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; 
    font-size: 11pt; 
    line-height: 1.6; 
    color: #333; 
}
.content { max-width: 100%; }
h1, h2, h3 { color: #000; margin-top: 1em; margin-bottom: 0.5em; }
h1 { font-size: 20pt; }
h2 { font-size: 16pt; }
h3 { font-size: 13pt; }
pre, code { 
    background: #f5f5f7; 
    padding: 8px; 
    border-radius: 4px; 
    font-family: 'Consolas', monospace; 
    font-size: 10pt; 
}
blockquote { 
    border-left: 4px solid #ddd; 
    margin: 1em 0; 
    padding: 0.5em 1em; 
    background: #fafafa; 
}
table { border-collapse: collapse; width: 100%; margin: 1em 0; }
th, td { border: 1px solid #ddd; padding: 6px 10px; text-align: left; }
th { background: #f0f0f0; font-weight: bold; }
"""

DARK_PDF_CSS = """
@page { size: A4; margin: 2cm; }
body { 
    background: #1a1a1a; 
    font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; 
    font-size: 11pt; 
    line-height: 1.6; 
    color: #e0e0e0; 
}
.content { max-width: 100%; }
h1, h2, h3 { color: #fff; margin-top: 1em; margin-bottom: 0.5em; }
h1 { font-size: 20pt; }
h2 { font-size: 16pt; }
h3 { font-size: 13pt; }
pre, code { 
    background: #2a2a2a; 
    color: #e0e0e0; 
    padding: 8px; 
    border-radius: 4px; 
    font-family: 'Consolas', monospace; 
    font-size: 10pt; 
}
blockquote { 
    border-left: 4px solid #555; 
    margin: 1em 0; 
    padding: 0.5em 1em; 
    background: #252525; 
}
table { border-collapse: collapse; width: 100%; margin: 1em 0; }
th, td { border: 1px solid #444; padding: 6px 10px; text-align: left; }
th { background: #333; font-weight: bold; }
"""
